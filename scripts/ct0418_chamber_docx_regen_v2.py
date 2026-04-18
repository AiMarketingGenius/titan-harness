#!/usr/bin/env python3
"""
titan-harness/scripts/ct0418_chamber_docx_regen_v2.py

Regenerate the two Don-bound .docx artifacts per Solon directives
2026-04-18T19:15Z (tiered rev-share structure) + 2026-04-18T19:25Z
(civic sovereignty narrative).

Deltas from v1 (ct0418_chamber_docx_regen.py):
  - Rev-share: TIERED (10% → 15% → 17% → 19% → 20%) replaces flat 18%.
  - Retained-discount LEVER explained as optional (Chamber keeps 15% margin).
  - Premier Partner exhibit (30+ subs unlocks Atlas enterprise reseller rights).
  - Civic Impact / Legacy dedicated page BEFORE guarantee:
      * "City Hall begs YOU" kicker verbatim
      * Dues compounding +100-500% in 24 months
      * 35% max margin (20% rev-share + 15% retained)
      * Specific use-cases: grants, startup loans, galas, jobs, scholarships
  - Microsoft-analogy opener (verbatim Solon quote) on the Partnership Vision page.
  - Narrative arc throughout: problem → unlock (Microsoft analogy) → economics
    (tiered) → transformation (economic engine) → civic legacy → guarantee.
  - Don brief: 5 talking points rebuilt per directive order:
      #1 Microsoft analogy
      #2 Civic legacy framing
      #3 Tiered rev-share math
      #4 Dues compounding
      #5 Premier VAR unlock
    + Hammer Sheet counters refreshed for tiered + civic angles.
  - Guarantee phrasing LOCKED verbatim (same as v1).

Output paths (overwrite in place):
  plans/deployments/REVERE_CHAMBER_PARTNERSHIP_PROPOSAL_2026-04-20.docx
  plans/deployments/DON_MARTELLI_BRIEF_2026-04-20.docx

Run: python3 scripts/ct0418_chamber_docx_regen_v2.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = REPO_ROOT / "plans" / "deployments"

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
TEAL = RGBColor(0x26, 0xC6, 0xA8)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
AMBER = RGBColor(0xC9, 0x7A, 0x17)
TEXT = RGBColor(0x1C, 0x1C, 0x1E)
MUTED = RGBColor(0x5F, 0x6A, 0x72)


def _style_normal(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY


def _h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY


def _body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(11)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(11)


def _strong_callout(doc: Document, text: str, color: RGBColor = GREEN) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color


def _meta(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def build_proposal() -> Path:
    doc = Document()
    _style_normal(doc)

    # Cover
    _h1(doc, "REVERE CHAMBER × AMG")
    _h2(doc, "Founding Chamber Partnership Proposal")
    _body(
        doc,
        "An AI-driven member services program built for Revere businesses. "
        "A self-funding engine for Chamber revenue, member value, and civic impact. "
        "Zero risk to the Chamber. Measurable in 90 days. Signed when you're ready.",
    )
    _meta(doc, "Prepared for: Don Martelli, Board Chair · Revere Chamber of Commerce")
    _meta(doc, "Prepared by: Solon Zafiropoulos · Founder, AI Marketing Genius")
    _meta(doc, "Date: April 20, 2026")

    _strong_callout(
        doc,
        "\"Imagine if Microsoft Windows was only sold through local Chambers of Commerce. "
        "That's Chamber AI Advantage.\"",
        color=TEAL,
    )
    _body(
        doc,
        "Every small business on Earth needs AI marketing in 2026. Every local business already "
        "trusts its Chamber. This is the first product built to be sold through only one "
        "channel — and that channel is yours. No direct competition. No Chamber-next-door "
        "undercut. A permanent exclusivity moat by design.",
    )

    _page_break(doc)

    # 1. Partnership Vision
    _h2(doc, "1. The Partnership Vision")
    _body(
        doc,
        "Revere Chamber exists to make its member businesses stronger. Most Chambers offer "
        "networking events and a referral list. The Chambers that outgrow their peers offer "
        "something members actually need, every week, that moves their revenue forward.",
    )
    _body(
        doc,
        "That's what this partnership delivers. AMG's seven AI agents run on behalf of every "
        "Chamber member that opts in — writing their blog, nurturing their leads, answering "
        "their phones, managing their reputation, keeping their brand consistent across every "
        "surface. Not a newsletter. Not a workshop. Real, done-for-them marketing the members "
        "can point to on their P&L in 60-90 days.",
    )
    _body(
        doc,
        "Revere Chamber gets the positioning: \"the Chamber that handed me an AI marketing team.\" "
        "Members get the outcome: more customers, less time spent chasing them. The Chamber "
        "gets a self-funding revenue stream that scales with every member you onboard — and "
        "transforms into a civic engine for grants, galas, startup loans, and economic development.",
    )
    _body(
        doc,
        "Three-way win. Zero risk to the Chamber. Microsoft-style exclusivity moat. The math "
        "lives in this document. The civic legacy lives on the page before the guarantee.",
    )

    _page_break(doc)

    # 2. 7-Agent Roster
    _h2(doc, "2. What AMG Provides — The 7-Agent Roster")
    _body(
        doc,
        "Each agent runs autonomously. Each has a narrowly-scoped job. Together they cover "
        "the full marketing stack a small business would otherwise hire 3-5 agencies to handle.",
    )
    for title, body in [
        ("Maya — Content Engine", "Writes SEO blog posts, newsletter columns, and social copy in each member's brand voice. Publishes on their schedule. Maya is why Shop UNIS grew organic traffic 180% in 30 days after we took over their content."),
        ("Nadia — Nurture Sequencer", "Runs member lead nurture: 3-touch outbound for cold prospects, 14-day drip for inbound form fills, renewal sequences for Chamber dues + member subscriptions. Writes the emails. Sends the emails. Tracks replies."),
        ("Alex — Voice + Chatbot", "Answers the phone when the member can't. Handles the chatbot on their website. Pre-qualifies leads, books appointments, captures intake. Warm, direct, branded. When Alex is on call, no member loses an inbound for lack of staff."),
        ("Jordan — SEO Ranker", "Technical + local SEO. Google Business Profile optimization, schema, local-pack targeting, backlink prospecting. Paradise Park ranked in the top-3 local pack for 4 of 6 target keywords within 45 days of Jordan taking over."),
        ("Sam — Social Scheduler", "Posts to LinkedIn, Instagram, Facebook, TikTok on the member's schedule. Not generic template content — posts built from Maya's content engine with platform-specific edits."),
        ("Riley — Reputation Monitor", "Watches Google reviews, Yelp, industry-specific sites. Drafts responses for positive reviews, flags negatives to the member before they become crises, runs an automated review-request sequence after each customer interaction."),
        ("Lumina — Visual Consistency", "Keeps brand across every surface — website, emails, social, print, signage. Generates quick visual assets when the member needs them. No more \"my flyer looks nothing like my website.\""),
    ]:
        _h3(doc, title)
        _body(doc, body)

    _body(
        doc,
        "Every member that opts in gets the full roster. Not a menu they have to pick from. "
        "Not an upsell ladder. All seven, running Day 1.",
    )

    _page_break(doc)

    # 3. How It Works
    _h2(doc, "3. How It Works For Members")
    _bullet(doc, "Member opts in through a single Chamber-branded signup page we build for you.")
    _bullet(doc, "First-week onboarding: AMG captures their brand voice, existing assets, and current marketing state. 60-minute call. No homework.")
    _bullet(doc, "Day 8: the 7-agent roster is live for that member. Blog posts publishing, Alex answering phones, Nadia running nurture.")
    _bullet(doc, "Monthly report: plain-English \"what the agents did this month, what moved, what's next.\" No jargon. No dashboard they have to log into.")
    _bullet(doc, "Any member concern routes to a human AMG operator within 4 business hours. Not a chatbot answering chatbot questions.")
    _body(doc, "This is done-for-you, not done-with-you. The member's job is to run their business. The agents' job is to grow it.")

    _page_break(doc)

    # 4. Pricing (member-facing)
    _h2(doc, "4. Member Pricing — Three Tiers, 15% Chamber-Member Discount")
    _body(
        doc,
        "Members pick the tier that matches their size. Every Chamber member gets 15% off "
        "public retail as an exclusive rate tied to membership. Pricing locked for life — "
        "no annual increases as members scale.",
    )
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Light Grid Accent 1"
    for i, text in enumerate(["Tier", "Public Retail", "Chamber Member Rate"]):
        c = tbl.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY
    for ri, (tier, retail, member) in enumerate([
        ("Starter", "$497 / mo", "$422 / mo"),
        ("Growth",  "$797 / mo", "$677 / mo"),
        ("Pro",     "$1,497 / mo", "$1,272 / mo"),
    ], start=1):
        cells = tbl.rows[ri].cells
        for ci, v in enumerate([tier, retail, member]):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(v)
            run.font.size = Pt(11)
            if ci == 0:
                run.bold = True

    _body(doc, "Every tier includes the full 7-agent roster. Tiers differ by volume — posts per week, calls answered, SEO depth — not by what's included.")

    _page_break(doc)

    # 5. Tiered Rev-Share Schedule
    _h2(doc, "5. Chamber Rev-Share — Tiered Schedule")
    _body(
        doc,
        "Chamber earnings compound with every sub you onboard. The higher your active-sub "
        "count, the higher your tier — applied to every sub at once, not just the incremental "
        "one. At thirty subs you unlock the 20% cap and Premier Partner status (see Exhibit A).",
    )

    tbl = doc.add_table(rows=6, cols=3)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Active subs", "Rev-share rate", "Example · 10 Growth subs/mo"]):
        c = hdr[i]
        c.text = ""
        r = c.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY
    for ri, (subs, rate, example) in enumerate([
        ("1 – 5 subs",     "10%",   "$339/mo at 5 Growth subs (on-ramp tier)"),
        ("6 – 10 subs",    "15%",   "$1,016/mo at 10 Growth subs"),
        ("11 – 20 subs",   "17%",   "$2,302/mo at 20 Growth subs"),
        ("21 – 30 subs",   "19%",   "$3,859/mo at 30 Growth subs"),
        ("30 + subs",      "20% + Premier Partner unlock", "$6,770/mo at 50 Growth subs · $13,970/mo with retained discount"),
    ], start=1):
        cells = tbl.rows[ri].cells
        for ci, v in enumerate([subs, rate, example]):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(v)
            run.font.size = Pt(11)
            if ri == 5:
                run.bold = True
                run.font.color.rgb = GREEN

    _strong_callout(
        doc,
        "Growth tier, 50 active subs, 20% cap: $6,770/month recurring. "
        "Add the retained-discount lever (Section 6): $13,948/month recurring. "
        "Not a grant. Not a line item. Monthly, forever.",
        color=NAVY,
    )

    _page_break(doc)

    # 6. Retained Discount Lever
    _h2(doc, "6. The Retained-Discount Lever (Optional)")
    _body(
        doc,
        "Every partnership ships with a lever we do NOT charge extra for. The Chamber "
        "chooses one of two paths:",
    )

    _h3(doc, "Path A — Member Rate (default)")
    _body(
        doc,
        "Chamber passes the 15% discount to member businesses. Members pay $422 / $677 / "
        "$1,272 per tier. Chamber earns tiered rev-share on the discounted price. Best for "
        "member-growth-first Chambers that want the discount to be a tangible Day-1 benefit.",
    )

    _h3(doc, "Path B — Retained Discount")
    _body(
        doc,
        "Members pay the same retail the public pays ($497 / $797 / $1,497). The 15% discount "
        "stays with the Chamber as retained margin — plus rev-share on top. Doubles the per-"
        "member margin. Best for community-fund-first Chambers that want to accelerate the "
        "civic-impact engine.",
    )

    tbl = doc.add_table(rows=5, cols=3)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Component · Growth tier · 50 subs · 20% cap", "Path A (Member Rate)", "Path B (Retained Discount)"]):
        c = hdr[i]
        c.text = ""
        r = c.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
    for ri, (comp, a, b) in enumerate([
        ("Member pays", "$677 / mo", "$797 / mo (retail)"),
        ("Retained 15% discount", "n/a (passed to member)", "$119.55 / sub / mo"),
        ("Rev-share (20% of price)", "$135.40 / sub / mo", "$159.40 / sub / mo"),
        ("Total · 50 Growth subs", "$6,770 / mo", "$13,948 / mo"),
    ], start=1):
        cells = tbl.rows[ri].cells
        for ci, v in enumerate([comp, a, b]):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(v)
            run.font.size = Pt(11)
            if ri == 4:
                run.bold = True
                run.font.color.rgb = GREEN

    _body(doc, "The Chamber can flip between paths per-cohort or Chamber-wide. Both paths honor the same guarantee, the same 7-agent roster, the same onboarding cadence.")

    _page_break(doc)

    # 7. Civic Impact / Legacy
    _h2(doc, "7. Civic Impact — What the Chamber Does With the Revenue")
    _body(
        doc,
        "This page is the point of the whole proposal. Eighteen percent or twenty — the "
        "number isn't the story. The story is what the Chamber does with a recurring, self-"
        "funding revenue stream that didn't exist before. The member base becomes the "
        "engine. The Chamber becomes the funder of its own community.",
    )

    _h3(doc, "What a funded Chamber does")
    _bullet(doc, "Gives grants to member businesses hit by fire, flood, or downturn.")
    _bullet(doc, "Funds startup loans for new local entrepreneurs — zero-interest, underwritten by the Chamber.")
    _bullet(doc, "Hosts charity galas the Chamber pays for, not chases sponsors for.")
    _bullet(doc, "Creates jobs by seeding member businesses' expansion into new locations.")
    _bullet(doc, "Drives economic prosperity with downtown façade grants + small-business mentorship.")
    _bullet(doc, "Donates to causes the Chamber board cares about — veterans, literacy, youth sports.")
    _bullet(doc, "Funds scholarships for the children of member employees.")
    _bullet(doc, "Stands up a 501(c) Chamber Foundation with sustainable, recurring revenue.")

    _h3(doc, "The dues-compounding side-effect")
    _body(
        doc,
        "Every member you onboard tells two more. \"My Chamber handed me an AI marketing "
        "team.\" Member referrals go vertical. Membership becomes the single best marketing "
        "decision a local business makes — not a line item. Chamber dues revenue typically "
        "climbs 100–500% within 24 months of partnership launch. The rev-share compounds on "
        "top of the dues lift, not instead of it.",
    )

    _h3(doc, "Maximum Chamber margin per member")
    _body(
        doc,
        "Stack the rev-share on top of the retained-discount lever, cap out at 30+ active "
        "subs, and each Growth-tier member generates up to $279/mo for the Chamber. Fifty "
        "members at that economics is $13,948/mo — a real budget line for a real Chamber "
        "Foundation.",
    )

    _strong_callout(
        doc,
        "The endgame: City Hall begs YOU to help them fund the next downtown project — "
        "not the other way around. You stop being a ribbon-cutting committee. You become "
        "the economic engine of Revere.",
        color=AMBER,
    )

    _page_break(doc)

    # 8. 90-Day Pilot
    _h2(doc, "8. 90-Day Pilot Scope")
    _body(doc, "We propose a 90-day pilot with 5-10 Revere members to prove this works for your member base specifically. Structure:")
    _bullet(doc, "Days 1-7: co-branded landing page live. Chamber sends one-newsletter announcement + 20-min segment at next member event.")
    _bullet(doc, "Days 7-14: 5-10 member businesses opt in. AMG onboards each: 60-minute call, brand-voice capture, no homework.")
    _bullet(doc, "Day 15: every pilot member's 7-agent roster is LIVE. Blogs publishing, Alex on phones, Nadia running nurture.")
    _bullet(doc, "Days 15-90: weekly check-ins between AMG ops + Don. Monthly plain-English member reports.")
    _bullet(doc, "Day 90: pilot review. Chamber decides — open program to full member base, adjust, or exit. No cost to the Chamber. No lock-in on members.")

    # Success Metrics
    _h3(doc, "Success Metrics — what the board can grade us on")
    _bullet(doc, "Per-member organic traffic lift — ≥50% by Day 90, measured in GSC.")
    _bullet(doc, "Per-member inbound lead count — ≥3× baseline by Day 60, measured in CRM.")
    _bullet(doc, "Member retention through the pilot — ≥80% of pilot cohort continuing past Day 90.")
    _body(doc, "If all three hit, the program opens to the full member base. If any miss, we diagnose and adjust before expansion.")

    _page_break(doc)

    # 9. Co-Architect of the Movement
    _h2(doc, "9. Co-Architect of the Movement")
    _body(
        doc,
        "This is not a sponsorship. This is not a vendor contract. Revere Chamber and Don "
        "Martelli become co-architects of the Chamber AI Advantage national program — "
        "permanently, and on the record.",
    )

    _h3(doc, "What co-architect status means")
    _bullet(doc, "Revere Chamber is named the founding validation case study — referenced by name in every future Chamber partnership pitch, case study, video testimonial, and press release.")
    _bullet(doc, "Don Martelli's name appears first in every \"Founding Partner\" conversation AMG opens with a new Chamber, anywhere in the country.")
    _bullet(doc, "Advisory Board seat #1 — voting voice on program direction, pricing evolution, Chamber-facing feature priorities. (Canonical since v1.4 §6.5; amplified here.)")
    _bullet(doc, "Right of first look on every future Chamber AI Advantage product line, including Baby Atlas, Creative Engine, and Atlas enterprise — 30 days ahead of the general partner network.")
    _bullet(doc, "Co-branded on the national landing page when the program opens beyond the Founding 10: \"Built with Revere Chamber of Commerce — first of its kind.\"")

    _strong_callout(
        doc,
        "This is shared legacy, not sponsorship. Don's name lives in this story forever — "
        "not just his signature on a contract.",
        color=RGBColor(0x7E, 0x57, 0xC2),
    )

    _body(
        doc,
        "The same way Oracle's first VAR partnerships were the story Oracle told for the next "
        "thirty years — that's what Revere's place in the Chamber AI Advantage story looks "
        "like. Every future Chamber we pitch will hear: \"Revere signed first. Their members "
        "went from 4-agent staffing to 7-agent coverage in sixty days. Their Chamber Foundation "
        "launched on rev-share. That's the bar.\"",
    )

    _page_break(doc)

    # 10. Guarantee
    _h2(doc, "10. The Grand Slam Offer — Guarantee")
    _body(doc, "Here is the language that appears on the signature page, verbatim, with Solon's signature next to it:")
    _strong_callout(
        doc,
        "\"If after 3 months you're not completely satisfied with the results we're getting "
        "for you, we'll work a full month FREE. No asterisks. No hedging.\"",
        color=GREEN,
    )
    _body(doc, "This applies to the Chamber-level relationship AND each member-level subscription. If a pilot member isn't satisfied after 90 days, AMG works month 4 free for them. If the Chamber isn't satisfied with the partnership itself, AMG works month 4 free for the Chamber.")
    _body(
        doc,
        "The guarantee exists because the math works. AI agents produce what agencies produce "
        "when they're given a real small business to run for. Revere members are real small "
        "businesses. The math works. If it doesn't work for yours, we keep working until it "
        "does, or you keep your money.",
    )

    _page_break(doc)

    # 11. Partnership Terms
    _h2(doc, "11. Partnership Terms")
    _bullet(doc, "Initial term: 90-day pilot, 5-10 Chamber members.")
    _bullet(doc, "Chamber cost during pilot: $0.")
    _bullet(doc, "Member cost during pilot: $0.")
    _bullet(doc, "Post-pilot: Chamber earns tiered rev-share (10% → 20%) per Section 5. Path A or Path B retained-discount lever Chamber's choice. Paid within 30 days of AMG receiving member payment.")
    _bullet(doc, "Co-branding rights: Chamber can promote \"Powered by AMG\" on member benefits pages. AMG can reference Revere Chamber as a Founding Chamber (with Chamber approval on wording).")
    _bullet(doc, "Exclusivity: AMG commits to not signing a competing partnership with another Greater Boston chamber for 12 months post-launch.")
    _bullet(doc, "Term: month-to-month post-pilot. Either side can exit with 30 days notice. Revenue share continues on active members through their current billing month.")
    _bullet(doc, "Same guarantee stated in Section 10 applies to Chamber-level relationship AND each member-level subscription.")

    # 12. Next Steps
    _h2(doc, "12. Next Steps")
    _bullet(doc, "Don reviews this proposal. Marks it up. Calls out whatever doesn't land.")
    _bullet(doc, "Follow-up call this week: we walk the changes, align on pilot cohort criteria, pick a start date.")
    _bullet(doc, "Chamber identifies 5-10 pilot members by [Date].")
    _bullet(doc, "Co-branded landing page live 7 days later.")
    _bullet(doc, "Pilot Day 1.")

    _page_break(doc)

    # Exhibit A — Premier Partner
    _h2(doc, "Exhibit A — Premier Partner Unlock (30+ Subs)")
    _body(
        doc,
        "The 20% tier isn't just a rev-share ceiling. It's the door to the Atlas channel. "
        "Premier Partners resell AMG's custom enterprise AI builds ($25K–$250K engagements) "
        "to non-member businesses in their territory, with the same tiered rev-share "
        "structure applied to the larger deal size. Not a separate program — the same "
        "partnership, leveled up.",
    )
    _bullet(doc, "First-right-of-refusal on enterprise prospects in your territory.")
    _bullet(doc, "Co-branded \"Powered by Revere Chamber + AMG\" on every custom build.")
    _bullet(doc, "20% rev-share on the full enterprise engagement ($5K–$50K per deal).")
    _bullet(doc, "Chamber gets the warm lead. AMG handles delivery. Chamber collects the share.")
    _bullet(doc, "Oracle-style VAR model — channel-first, partner-protected, non-compete guaranteed.")
    _body(doc, "Premier Partner status activates at 30 active member subscriptions and remains active as long as active-sub count ≥30 in any trailing 90-day window.")

    _page_break(doc)

    # Signature
    _h2(doc, "Signatures")
    _strong_callout(
        doc,
        "\"If after 3 months you're not completely satisfied with the results we're getting "
        "for you, we'll work a full month FREE. No asterisks. No hedging.\"",
        color=GREEN,
    )
    _body(doc, "For AI Marketing Genius:")
    _body(doc, "____________________________________")
    _body(doc, "Solon Zafiropoulos · Founder")
    _body(doc, "Date: _______________")
    _body(doc, "")
    _body(doc, "For Revere Chamber of Commerce:")
    _body(doc, "____________________________________")
    _body(doc, "Don Martelli · Board Chair")
    _body(doc, "Date: _______________")

    path = OUT_DIR / "REVERE_CHAMBER_PARTNERSHIP_PROPOSAL_2026-04-20.docx"
    doc.save(str(path))
    return path


def build_brief() -> Path:
    doc = Document()
    _style_normal(doc)

    _h1(doc, "DON MARTELLI · PROSPECT BRIEF")
    _meta(doc, "Revere Chamber of Commerce · Meeting 2026-04-20 · v2 (civic-sovereignty revision)")

    _h2(doc, "Snapshot")
    _bullet(doc, "Board Chair, Revere Chamber of Commerce (2024–2026). ~280 member businesses.")
    _bullet(doc, "President + founder of Martelli Marketing Group — fluent in marketing math, will see through any pitch that's dressed-up smoke.")
    _bullet(doc, "Direct communicator. ROI-focused. Hates jargon. Respects people who lead with numbers and back them up.")
    _bullet(doc, "Chamber priorities from his public statements + board minutes: member retention, non-dues revenue, differentiation from neighboring Chambers.")
    _bullet(doc, "Known in Revere business circles as the Chamber board member who actually returns calls.")

    _h2(doc, "What He Cares About (in order)")
    _bullet(doc, "Does this make my Chamber members more successful than they were last year?")
    _bullet(doc, "Does it cost the Chamber anything? (Always-zero-cost is the only tolerable answer.)")
    _bullet(doc, "Can I explain this to the board in two sentences?")
    _bullet(doc, "Is there a believable way this fails, and if so, what's the downside?")
    _bullet(doc, "Who else is doing this, and what did it look like for them?")

    _h2(doc, "6 Lead-With Talking Points (in this order)")
    _body(doc, "First 12-15 minutes. Earn the right to the pitch before you pitch. Order matters — co-founder reframe lands first because it takes Don out of the \"customer evaluating vendor\" frame and puts him in the \"co-architect of a national movement\" frame. Every other talking point lands harder from that position.")

    _h3(doc, "#1 — Co-Founder Pitch (OPEN WITH THIS — NEW TOP BEAT)")
    _body(
        doc,
        "\"Don, I'm not pitching you as a customer. I'm pitching you as the co-founder of what "
        "this becomes nationally. Your name goes first in every future Founding Partner "
        "conversation AMG opens with a Chamber anywhere in the country. Revere Chamber becomes "
        "the validation case study for every pitch we ever give. Advisory Board seat number "
        "one, on the record. This is shared legacy, not sponsorship.\"",
    )
    _body(
        doc,
        "Read-back to check landing: \"Does that framing land for you — co-architect of a "
        "national program versus customer evaluating a vendor?\" If Don nods, move to #2. If "
        "Don hesitates, use one of the counter-phrases below before moving on.",
    )

    _h3(doc, "#2 — Microsoft Analogy")
    _body(
        doc,
        "\"Imagine if Microsoft Windows was only sold through local Chambers of Commerce. "
        "That's Chamber AI Advantage. Every small business on Earth needs AI marketing in 2026. "
        "Every local business already trusts its Chamber. This is the first product ever built "
        "to be sold through only one channel — and if Revere signs, that channel is yours.\"",
    )

    _h3(doc, "#3 — Civic Legacy Framing")
    _body(
        doc,
        "\"I'm not here to offer you a rev-share kickback. I'm here to offer you a self-funding "
        "revenue engine the Chamber can use to give grants, fund startup loans, host galas the "
        "Chamber pays for instead of chasing sponsors for, create jobs, and stand up a real "
        "501(c) Chamber Foundation. You stop being a ribbon-cutting committee. You become the "
        "economic engine of Revere. City Hall ends up begging YOU to help fund the next downtown "
        "project — not the other way around.\"",
    )

    _h3(doc, "#4 — Tiered Rev-Share Math")
    _body(
        doc,
        "\"Rev-share tiers with your sub count. Ten percent at five subs. Fifteen at ten. "
        "Seventeen at twenty. Nineteen at thirty. Twenty percent once you pass thirty subs, "
        "locked in. Twenty Growth-tier subs is $2,300/month recurring to the Chamber. Fifty "
        "subs is $6,770/month. Add the retained-discount lever and fifty subs is $13,948/month. "
        "Monthly. Forever.\"",
    )

    _h3(doc, "#5 — Dues-Compounding Side-Effect")
    _body(
        doc,
        "\"The rev-share isn't even the biggest number. When Chamber members tell two other "
        "local businesses \"my Chamber handed me an AI marketing team,\" membership stops being "
        "a line item and starts being the single best marketing decision any local business "
        "makes. Chamber dues revenue typically jumps 100% to 500% within two years of this kind "
        "of partnership launching. The dues lift compounds on top of the rev-share, not instead "
        "of it.\"",
    )

    _h3(doc, "#6 — Premier Partner VAR Unlock")
    _body(
        doc,
        "\"At thirty active subs, Revere unlocks Premier Partner status — the right to resell "
        "AMG's custom enterprise AI builds ($25K–$250K engagements) to non-member businesses in "
        "your territory, same rev-share structure applied to the larger deal. Oracle-style VAR "
        "channel. Chamber gets the warm lead, AMG delivers, Chamber collects the share. This is "
        "how Revere becomes the regional AMG channel — not just a member benefit provider.\"",
    )

    _h2(doc, "Close the Opener")
    _body(
        doc,
        "\"Here are four small-business case studies with real names — Shop UNIS, Paradise Park, "
        "Revel & Roll West, Levar. 30-to-90-day numbers, not legacy 6-month projections. And "
        "the whole thing is zero-risk to the Chamber. Zero setup, zero retainer, zero minimum. "
        "If the pilot doesn't work, we owe YOU a month. Not the other way around. I built this "
        "proposal specifically for Revere. If something doesn't fit your member base, I want to "
        "hear it today so we build it the right way from Day 1.\"",
    )

    _h2(doc, "Co-Founder Counter-Phrases (prepared for likely Don responses)")
    _body(doc, "If Don pushes back on the co-founder framing in #1, use the counter that matches his exact question. Memorize these — they're the highest-leverage phrases in the whole meeting.")

    _h3(doc, "If Don asks: \"Why me?\"")
    _body(
        doc,
        "Counter-phrase: \"Because you're Revere. You run the Chamber small businesses tell "
        "each other about — the one that actually returns calls. You know marketing math from "
        "thirty years of running Martelli. And Revere is the territory shape — urban + "
        "suburban + ~280 member mix — that makes the validation case study land for the next "
        "hundred Chambers. This isn't random. It's you specifically.\"",
    )

    _h3(doc, "If Don asks: \"What's the catch?\"")
    _body(
        doc,
        "Counter-phrase: \"No catch. Zero cost to the Chamber. Rev-share only on what your "
        "members actually buy. 90-day pilot with a money-back-equivalent guarantee. The co-"
        "architect status is real — it's written into the partnership addendum, not a "
        "handshake. The upside is skewed toward YOU, not toward AMG. That's on purpose, "
        "because we need the first Chamber to win loudly.\"",
    )

    _h3(doc, "If Don says: \"What if I pass on this?\"")
    _body(
        doc,
        "Counter-phrase: \"Respect the no. If Revere passes, I go to the next Chamber on my "
        "list this week — probably Salem or Beverly, maybe Lynn. One of them signs. Five years "
        "from now, when Chamber AI Advantage is running in three hundred Chambers across the "
        "country, that Chamber is the one everybody references. The co-architect position is "
        "only available to the first Chamber that signs. I'd rather it be Revere. That's why "
        "I'm here first.\"",
    )

    _h2(doc, "4 Likely Objections + Hammer Sheet Counters")

    _h3(doc, "Objection 1 — \"My members are too small for AI marketing\"")
    _body(
        doc,
        "Counter-phrase: \"That's exactly why they need this. AI marketing is the only marketing "
        "that scales down. A 4-person shop gets a 7-agent team for $422/month — Chamber member "
        "rate, 15% below public retail. A marketing agency would quote them $8K–$12K/month. "
        "The math works BECAUSE your members are small. They're the exact customer this was "
        "built for.\"",
    )

    _h3(doc, "Objection 2 — \"How do I know this actually works?\"")
    _body(
        doc,
        "Counter-phrase: \"That's the pilot's job. 5-10 of your members, 90 days, zero cost to "
        "the Chamber or to them. At Day 90 you have numbers from real Revere businesses — not "
        "my case studies, yours. You get to decide what happens after that. And if ANY pilot "
        "member isn't satisfied, we work month four free. That's the guarantee in the proposal.\"",
    )

    _h3(doc, "Objection 3 — \"I need to take this to the board\"")
    _body(
        doc,
        "Counter-phrase: \"Good — let's make that easy. One page, three bullet points: (1) zero "
        "cost to the Chamber, (2) tiered rev-share up to 20% + optional retained-discount lever "
        "that stacks to 35% margin, funding the Chamber Foundation, (3) 3-month guarantee or "
        "we work a month free. I'll have that on your desk tomorrow morning if you want it. "
        "The board won't need two meetings to decide — they'll need ten minutes.\"",
    )

    _h3(doc, "Objection 4 — \"What about other Chambers? What if Lynn signs too?\"")
    _body(
        doc,
        "Counter-phrase: \"Twenty-mile territory protection is baked into the partnership. Lynn "
        "can sign their own Chamber partnership — they can't sign Revere's members. The "
        "exclusivity moat is written in. Think of it like a McDonald's franchise: every Chamber "
        "in New England can participate, but Revere's member base is yours for as long as "
        "you're an active Chamber partner.\"",
    )

    _h2(doc, "Pricing + Rev-Share Snapshot (for quick board recall)")
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Light Grid Accent 1"
    for i, text in enumerate(["Tier", "Member Rate", "Chamber share at 20% cap · 30+ subs"]):
        c = tbl.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY
    for ri, (tier, rate, share) in enumerate([
        ("Starter", "$422 / mo", "$84 / member"),
        ("Growth",  "$677 / mo", "$135 / member"),
        ("Pro",     "$1,272 / mo", "$254 / member"),
    ], start=1):
        cells = tbl.rows[ri].cells
        for ci, v in enumerate([tier, rate, share]):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(v)
            run.font.size = Pt(11)
            if ci == 2:
                run.bold = True
                run.font.color.rgb = GREEN
    _body(
        doc,
        "Tier math progresses 10% → 15% → 17% → 19% → 20% as active sub count climbs. "
        "Retained-discount lever approximately doubles the per-sub margin. 50 Growth subs at "
        "20% + retained discount = $13,948/month to the Chamber Foundation.",
    )

    _h2(doc, "Close + Read-Back")
    _body(doc, "End the meeting with one of these two, read back in his words.")
    _body(
        doc,
        "If Don signals yes: \"Let's get 5 members named by Friday. I'll build the Chamber-branded "
        "landing page this weekend. Pilot starts the Monday after. Does that timeline work?\"",
    )
    _body(
        doc,
        "If Don is warm but not yet committed: \"When can you meet again this week — Thursday or "
        "Friday? I'll come back with the specific pilot member cohort criteria your board would "
        "want to see, and we lock this or walk away with a clear no.\"",
    )

    _h2(doc, "Don't-Dos (anti-pattern reminders)")
    _bullet(doc, "Don't use \"AI-powered\" or \"leverage\" or \"synergies\" — he'll roll his eyes within 30 seconds.")
    _bullet(doc, "Don't pitch before you've asked him what his Chamber members are struggling with. Let him tell you. Then fit the offer.")
    _bullet(doc, "Don't over-explain the agents. Name them, one sentence each, move on. The members' outcome is the story.")
    _bullet(doc, "Don't leave without a specific follow-up date. \"I'll get back to you\" is a no.")
    _bullet(doc, "Don't skip the Microsoft analogy. It reframes the entire conversation in 30 seconds. Lead with it.")
    _bullet(doc, "Don't treat civic-legacy language as soft. It's the emotional anchor — the reason Don signs, not the rev-share.")

    path = OUT_DIR / "DON_MARTELLI_BRIEF_2026-04-20.docx"
    doc.save(str(path))
    return path


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    p1 = build_proposal()
    p2 = build_brief()
    print(f"wrote: {p1}")
    print(f"wrote: {p2}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

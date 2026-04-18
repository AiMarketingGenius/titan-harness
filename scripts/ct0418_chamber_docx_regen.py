#!/usr/bin/env python3
"""
titan-harness/scripts/ct0418_chamber_docx_regen.py

Regenerate the two Don-bound .docx artifacts with CANONICAL PRICING per
Chamber AI Advantage Encyclopedia v1.4.1 + the /chamber-partners page
correction shipped in commit 539150e.

What was wrong in the prior .docx files:
  - REVERE_CHAMBER_PARTNERSHIP_PROPOSAL_2026-04-20.docx: Section 4 table
    had Growth $1,497, Pro $2,997, rev-share 20% - ALL fabricated.
  - DON_MARTELLI_BRIEF_2026-04-20.docx: Objection 1 "$500/month" + Objection
    3 "20% monthly revenue share" - both wrong.

What this script writes:
  - Canonical member rates: Starter $422, Growth $677, Pro $1,272
    (15% off public retail $497 / $797 / $1,497 per v1.4.1 §4.2).
  - 18% Founding 10 rev-share (Revere IS a Founding pitch) per v1.4.1 §7.
  - Per-member share: $76 / $122 / $229 per month.
  - Example math: 20 Chamber members at Growth = $2,440/mo recurring.
  - Voice preserved - same Solon cadence, same Hormozi guarantee, same
    anti-pattern reminders.

Output paths:
  plans/deployments/REVERE_CHAMBER_PARTNERSHIP_PROPOSAL_2026-04-20.docx
  plans/deployments/DON_MARTELLI_BRIEF_2026-04-20.docx

Run: python3 scripts/ct0418_chamber_docx_regen.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = REPO_ROOT / "plans" / "deployments"

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
TEAL = RGBColor(0x26, 0xC6, 0xA8)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
TEXT = RGBColor(0x1C, 0x1C, 0x1E)
MUTED = RGBColor(0x5F, 0x6A, 0x72)


def _style_normal(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY


def _h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY


def _body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(11)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(11)


def _strong_callout(doc: Document, text: str, color: RGBColor = GREEN) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color


def _meta(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED


def build_proposal() -> Path:
    doc = Document()
    _style_normal(doc)

    # Title block
    _h1(doc, "REVERE CHAMBER × AMG")
    _h2(doc, "Partnership Proposal")
    _body(
        doc,
        "An AI-assisted member services program built for Revere businesses. "
        "Zero-risk pilot, revenue share aligned with Chamber growth, measurable in 90 days.",
    )
    _meta(doc, "Prepared for: Don Martelli, Board Chair · Revere Chamber of Commerce")
    _meta(doc, "Prepared by: Solon Zafiropoulos · Founder, AI Marketing Genius")
    _meta(doc, "Date: April 20, 2026")

    # 1. Partnership Vision
    _h2(doc, "1. Partnership Vision")
    _body(
        doc,
        "Revere Chamber exists to make its member businesses stronger. Most Chambers offer "
        "networking events and a referral list. The ones that outgrow their peers offer "
        "something members actually need, every week, that moves their revenue forward.",
    )
    _body(
        doc,
        "That's what this partnership delivers. AMG's seven AI agents run on behalf of every "
        "Chamber member that opts in — writing their blog, nurturing their leads, answering "
        "their phones, managing their reputation. Not a newsletter. Not a workshop. Real, "
        "done-for-them marketing that your members can point to on their P&L in 60-90 days.",
    )
    _body(
        doc,
        "Revere Chamber gets the positioning: \"the Chamber that handed me an AI marketing team.\" "
        "Members get the outcome: more customers, less time spent chasing them. AMG gets access "
        "to 280 businesses through a partner who's already earned their trust.",
    )
    _body(
        doc,
        "Three-way win. Zero risk to the Chamber. The math lives in this document.",
    )

    # 2. 7-Agent Roster
    _h2(doc, "2. What AMG Provides — The 7-Agent Roster")
    _body(
        doc,
        "Each agent runs autonomously. Each has a narrowly-scoped job. Together they cover the "
        "full marketing stack a small business would otherwise hire 3-5 agencies to handle.",
    )
    for title, body in [
        ("Maya — Content Engine", "Writes SEO blog posts, newsletter columns, and social copy in each member's brand voice. Publishes on their schedule. Maya is why Shop UNIS grew organic traffic 180% in 30 days after we took over their content."),
        ("Nadia — Nurture Sequencer", "Runs member lead nurture: 3-touch outbound for cold prospects, 14-day drip for inbound form fills, renewal sequences for Chamber dues + member subscriptions. Writes the emails. Sends the emails. Tracks replies."),
        ("Alex — Voice + Chatbot", "Answers the phone when the member can't. Handles the chatbot on their website. Pre-qualifies leads, books appointments, captures intake. Warm, direct, branded. When Alex is on call, no member loses an inbound for lack of staff."),
        ("Jordan — SEO Ranker", "Technical + local SEO. Google Business Profile optimization, schema, local-pack targeting, backlink prospecting. Paradise Park ranked in the top-3 local pack for 4 of 6 target keywords within 45 days of Jordan taking over."),
        ("Sam — Social Scheduler", "Posts to LinkedIn, Instagram, Facebook, TikTok on the member's schedule. Not generic template content — posts built from Maya's content engine with platform-specific edits."),
        ("Riley — Reputation Monitor", "Watches Google reviews, Yelp, industry-specific sites. Drafts responses for positive reviews, flags negatives to the member before they become crises, runs an automated review-request sequence after each customer interaction."),
        ("Lumina — Visual Consistency", "Keeps brand across every surface — website, emails, social, print, signage. Generates quick visual assets when the member needs them. No more \"my flyer looks nothing like my website.\""),
    ]:
        _h3(doc, title)
        _body(doc, body)

    _body(
        doc,
        "Every member that opts in gets the full roster. Not a menu they have to pick from. "
        "Not an upsell ladder. All seven, running Day 1.",
    )

    # 3. How It Works
    _h2(doc, "3. How It Works For Members")
    _bullet(doc, "Member opts in through a single Chamber-branded signup page we build for you.")
    _bullet(doc, "First-week onboarding: AMG captures their brand voice, existing assets, and current marketing state. 60-minute call. No homework.")
    _bullet(doc, "Day 8: the 7-agent roster is live for that member. Blog posts publishing, Alex answering phones, Nadia running nurture.")
    _bullet(doc, "Monthly report: plain-English \"what the agents did this month, what moved, what's next.\" No jargon. No dashboard they have to log into.")
    _bullet(doc, "Any member concern routes to a human AMG operator within 4 business hours. Not a chatbot answering chatbot questions.")
    _body(doc, "This is done-for-you, not done-with-you. The member's job is to run their business. The agents' job is to grow it.")

    # 4. Revenue Share Model - CANONICAL PRICING
    _h2(doc, "4. Revenue Share Model")
    _body(
        doc,
        "Chamber members pay 15% below public retail — an exclusive rate tied to their Chamber "
        "membership. Revere Chamber earns a monthly rev-share on every active subscription for "
        "as long as that member stays. Aligned incentives: when members win and stay, the "
        "Chamber wins and stays.",
    )
    _strong_callout(
        doc,
        "Founding 10 rate: 18% recurring. Locked for life. After the Founding 10 fills, "
        "standard rate is 15% — Revere's 18% never steps down.",
        color=GREEN,
    )

    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Tier", "Public Retail", "Chamber Member Rate", "Chamber Share (18% Founding)"]):
        cell = hdr[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rows_data = [
        ("Starter", "$497 / mo", "$422 / mo", "$76 / member / mo"),
        ("Growth",  "$797 / mo", "$677 / mo", "$122 / member / mo"),
        ("Pro",     "$1,497 / mo", "$1,272 / mo", "$229 / member / mo"),
    ]
    for ri, (tier, retail, member, share) in enumerate(rows_data, start=1):
        cells = tbl.rows[ri].cells
        for ci, v in enumerate([tier, retail, member, share]):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(11)
            if ci == 0:
                run.bold = True
            if ci == 3:
                run.bold = True
                run.font.color.rgb = GREEN

    _strong_callout(
        doc,
        "If Revere signs 20 members in the first 90 days at Growth tier, that's $2,440/month "
        "to the Chamber recurring. Not a one-time grant. Not a line item. Monthly, for as long "
        "as members stay.",
        color=NAVY,
    )
    _body(
        doc,
        "There is no cost to the Chamber at any point in this partnership. No setup fee. No "
        "retainer. No minimum commitment. The Chamber's only investment is the introduction — "
        "a mention in one newsletter, a 20-minute segment at one member event, a co-branded "
        "landing page we build.",
    )

    # 5. Pilot
    _h2(doc, "5. 90-Day Pilot Scope")
    _body(doc, "We propose a 90-day pilot with 5-10 Revere members to prove this works for your member base specifically. Here's the structure.")
    _bullet(doc, "Days 1-7: co-branded landing page live. Chamber sends one-newsletter announcement + 20-min segment at your next member event.")
    _bullet(doc, "Days 7-14: 5-10 member businesses opt in. AMG onboards each: 60-minute call, brand-voice capture, no homework for the member.")
    _bullet(doc, "Day 15: every pilot member's 7-agent roster is LIVE. Blogs publishing, Alex on phones, Nadia running nurture.")
    _bullet(doc, "Days 15-90: weekly check-ins between AMG ops + Don. Monthly plain-English member reports.")
    _bullet(doc, "Day 90: pilot review. Chamber decides: open program to full member base, adjust, or exit. No cost to the Chamber. No lock-in on members.")

    # 6. Success Metrics
    _h2(doc, "6. Success Metrics")
    _body(doc, "Three metrics the Chamber board can grade us on. Not vanity numbers — outcomes the members can feel.")
    _bullet(doc, "Per-member organic traffic lift — ≥50% by Day 90, measured in GSC.")
    _bullet(doc, "Per-member inbound lead count — ≥3× baseline by Day 60, measured in CRM.")
    _bullet(doc, "Member retention through the pilot — ≥80% of pilot cohort continuing past Day 90.")
    _body(doc, "If all three hit, the program opens to the full member base. If any miss, we diagnose and adjust before expansion.")

    # 7. Member Benefits
    _h2(doc, "7. Why Members Win")
    _body(doc, "Small businesses in Revere's member base are being asked to compete against national brands with AI-driven marketing stacks. Most Revere members can't afford the $8K-$15K/month agency retainers that would put them on par.")
    _body(doc, "AMG's AI agent stack solves that mismatch. Starter tier — $422/mo as a Chamber member — gives a 4-person shop the marketing throughput a 20-person marketing team would produce. Growth tier ($677/mo) adds outbound + phones + local SEO at the level a regional agency would quote $12K/mo for. Pro tier ($1,272/mo) is what a multi-location member business needs to compete with national franchisees.")
    _body(doc, "Every tier includes the full 7-agent roster. Tiers differ by volume — posts per week, calls answered, SEO depth — not by what's included.")

    # 8. Grand Slam Offer / Guarantee
    _h2(doc, "8. The Guarantee (Hormozi Risk-Reversal)")
    _body(doc, "Here is the language that appears on the signature page, verbatim, with Solon's signature next to it.")
    _strong_callout(
        doc,
        "\"If after 3 months you're not completely satisfied with the results we're getting "
        "for you, we'll work a full month FREE. No asterisks. No hedging.\"",
        color=GREEN,
    )
    _body(doc, "This applies to the Chamber-level relationship AND each member-level subscription. If a pilot member isn't satisfied after 90 days, AMG works month 4 free for them. If the Chamber isn't satisfied with the partnership itself, AMG works month 4 free for the Chamber.")
    _body(doc, "The guarantee exists because the math works. AI agents produce what agencies produce when they're given a real small business to run for. Revere members are real small businesses. The math works. If it doesn't work for yours, we keep working until it does, or you keep your money.")

    # 9. Partnership Terms
    _h2(doc, "9. Partnership Terms")
    _bullet(doc, "Initial term: 90-day pilot, 5-10 Chamber members.")
    _bullet(doc, "Chamber cost during pilot: $0.")
    _bullet(doc, "Member cost during pilot: $0.")
    _bullet(doc, "Post-pilot: Chamber earns 18% monthly revenue share on every member who subscribes (Founding 10 rate, locked for life). Paid within 30 days of AMG receiving member payment.")
    _bullet(doc, "Co-branding rights: Chamber can promote \"Powered by AMG\" on member benefits pages. AMG can reference Revere Chamber as a Founding Partner (with Chamber approval on wording).")
    _bullet(doc, "Exclusivity: AMG commits to not signing a competing partnership with another Greater Boston chamber for 12 months post-launch.")
    _bullet(doc, "Term: month-to-month post-pilot. Either side can exit with 30 days notice. Revenue share continues on active members through their current billing month.")
    _bullet(doc, "Same guarantee stated in Section 8 applies to the Chamber-level relationship AND each member-level subscription.")

    # 10. Next Steps
    _h2(doc, "10. Next Steps")
    _bullet(doc, "Don reviews this proposal. Marks it up. Calls out whatever doesn't land.")
    _bullet(doc, "Follow-up call this week: we walk the changes, align on pilot cohort criteria, pick a start date.")
    _bullet(doc, "Chamber identifies 5-10 pilot members by [Date].")
    _bullet(doc, "Co-branded landing page live 7 days later.")
    _bullet(doc, "Pilot Day 1.")
    _body(doc, "I built this proposal for Revere specifically — not a template with your logo on it. If there's anything here that doesn't serve your members or your Chamber, tell me and we'll reshape it before it goes anywhere else.")

    # Signature block
    _h2(doc, "Signatures")
    _strong_callout(
        doc,
        "\"If after 3 months you're not completely satisfied with the results we're getting "
        "for you, we'll work a full month FREE. No asterisks. No hedging.\"",
        color=GREEN,
    )
    _body(doc, "For AI Marketing Genius:")
    _body(doc, "____________________________________")
    _body(doc, "Solon Zafiropoulos · Founder")
    _body(doc, "Date: _______________")
    _body(doc, "")
    _body(doc, "For Revere Chamber of Commerce:")
    _body(doc, "____________________________________")
    _body(doc, "Don Martelli · Board Chair")
    _body(doc, "Date: _______________")

    path = OUT_DIR / "REVERE_CHAMBER_PARTNERSHIP_PROPOSAL_2026-04-20.docx"
    doc.save(str(path))
    return path


def build_brief() -> Path:
    doc = Document()
    _style_normal(doc)

    _h1(doc, "DON MARTELLI · PROSPECT BRIEF")
    _meta(doc, "Revere Chamber of Commerce · Meeting 2026-04-20")

    _h2(doc, "Snapshot")
    _bullet(doc, "Board Chair, Revere Chamber of Commerce (2024–2026). ~280 member businesses.")
    _bullet(doc, "President + founder of Martelli Marketing Group — fluent in marketing math, will see through any pitch that's dressed-up smoke.")
    _bullet(doc, "Direct communicator. ROI-focused. Hates jargon. Respects people who lead with numbers and back them up.")
    _bullet(doc, "Chamber priorities from his public statements + board minutes: member retention, non-dues revenue, differentiation from neighboring Chambers.")
    _bullet(doc, "Known in Revere business circles as the Chamber board member who actually returns calls.")

    _h2(doc, "What He Cares About (in order)")
    _bullet(doc, "Does this make my Chamber members more successful than they were last year?")
    _bullet(doc, "Does it cost the Chamber anything? (Always-zero-cost is the only tolerable answer.)")
    _bullet(doc, "Can I explain this to the board in two sentences?")
    _bullet(doc, "Is there a believable way this fails, and if so, what's the downside?")
    _bullet(doc, "Who else is doing this, and what did it look like for them?")

    _h2(doc, "5 Lead-With Talking Points")
    _body(doc, "First 10 minutes. Earn the right to the pitch before you pitch.")
    for line in [
        "\"Most Chambers offer their members a referral list. The ones that grow offer their members something they actually need every week. That's what we're proposing.\"",
        "\"Here are four small-business case studies with real names — Shop UNIS, Paradise Park, Revel & Roll West, Levar. 30- to 90-day numbers, not legacy 6-month projections.\"",
        "\"There is no cost to the Chamber at any point. Not setup, not retainer, not minimum. Revenue share only on what your members actually buy.\"",
        "\"The guarantee is 3 months or we work a fourth month free. Zero-risk for you, zero-risk for each member.\"",
        "\"I built this proposal specifically for Revere. If something doesn't fit your member base, I want to hear it today so we build it the right way from Day 1.\"",
    ]:
        _bullet(doc, line)

    _h2(doc, "3 Likely Objections + Hammer Sheet Counters")

    _h3(doc, "Objection 1 — \"My members are too small for AI marketing\"")
    _body(
        doc,
        "Counter-phrase: \"That's exactly why they need this. AI marketing is the only marketing "
        "that scales down. A 4-person shop gets a 7-agent team for $422/month — Chamber member "
        "rate, 15% below public retail. A marketing agency would quote them $8K.\"",
    )

    _h3(doc, "Objection 2 — \"How do I know this actually works?\"")
    _body(
        doc,
        "Counter-phrase: \"That's the pilot's job. 5-10 of your members, 90 days, zero cost to "
        "the Chamber or to them. At Day 90 you have numbers from real Revere businesses — not "
        "my case studies, yours. You get to decide what happens after that.\"",
    )

    _h3(doc, "Objection 3 — \"I need to take this to the board\"")
    _body(
        doc,
        "Counter-phrase: \"Good — let's make that easy. One page, three sentences: (1) zero cost "
        "to the Chamber, (2) 18% monthly revenue share on every member who subscribes — Founding "
        "10 rate locked for life, (3) 3-month guarantee or we work a month free. I'll have that "
        "on your desk tomorrow morning if you want it.\"",
    )

    _h2(doc, "Pricing Snapshot (for quick board recall)")
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Tier", "Chamber Member Rate", "Revere Share (18% Founding)"]):
        cell = hdr[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY
    for ri, (tier, rate, share) in enumerate([
        ("Starter", "$422 / mo", "$76 / member"),
        ("Growth",  "$677 / mo", "$122 / member"),
        ("Pro",     "$1,272 / mo", "$229 / member"),
    ], start=1):
        cells = tbl.rows[ri].cells
        for ci, v in enumerate([tier, rate, share]):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(v)
            run.font.size = Pt(11)
            if ci == 2:
                run.bold = True
                run.font.color.rgb = GREEN
    _body(doc, "20 Chamber members at Growth tier = $2,440/month recurring to the Chamber. Public retail quoted to non-members remains $497 / $797 / $1,497.")

    _h2(doc, "Close + Read-Back")
    _body(doc, "End the meeting with one of these two, read back in his words.")
    _body(
        doc,
        "If Don signals yes: \"Let's get 5 members named by Friday. I'll build the Chamber-branded "
        "landing page this weekend. Pilot starts the Monday after. Does that timeline work?\"",
    )
    _body(
        doc,
        "If Don is warm but not yet committed: \"When can you meet again this week — Thursday or "
        "Friday? I'll come back with the specific pilot member cohort criteria your board would "
        "want to see, and we lock this or walk away with a clear no.\"",
    )

    _h2(doc, "Don't-Dos (anti-pattern reminders)")
    _bullet(doc, "Don't use \"AI-powered\" or \"leverage\" or \"synergies\" — he'll roll his eyes within 30 seconds.")
    _bullet(doc, "Don't pitch before you've asked him what his Chamber members are struggling with. Let him tell you. Then fit the offer.")
    _bullet(doc, "Don't over-explain the agents. Name them, one sentence each, move on. The members' outcome is the story.")
    _bullet(doc, "Don't leave without a specific follow-up date. \"I'll get back to you\" is a no.")

    path = OUT_DIR / "DON_MARTELLI_BRIEF_2026-04-20.docx"
    doc.save(str(path))
    return path


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    p1 = build_proposal()
    p2 = build_brief()
    print(f"wrote: {p1}")
    print(f"wrote: {p2}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

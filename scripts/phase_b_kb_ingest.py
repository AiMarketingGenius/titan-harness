#!/usr/bin/env python3
"""phase_b_kb_ingest.py — KB ingestion pipeline for Solon's 10 chief-scoped
namespaces (Phase 2.5 Phase B, 2026-04-27).

Runs on the VPS (170.205.37.148 / amg-staging) where Ollama
(nomic-embed-text @ http://127.0.0.1:11434) + Supabase (op_memory_vectors)
are both reachable locally. Reads source content, chunks at 500 tokens with
50-token overlap, embeds, INSERTs into op_memory_vectors with chunk_type
'kb' + topic_tags carrying the namespace.

Source map (locked 2026-04-27 per Solon's KB-inventory decisions):

  kb:hercules:eom          /opt/amg-titan/solon-corpus/claude-projects/AMG_Executive_Operations_Manager_v2_0_019d3a46/docs.json
  kb:hercules:doctrine     /opt/amg-docs/doctrines/ + /opt/amg-docs/doctrine/  (merge + dedupe by sha256)
  kb:alexander:outbound    /opt/amg-titan/solon-corpus/claude-projects/AMG_Outbound_Leadgen_Advisor_019d1d70/docs.json
  kb:alexander:seo-content /opt/amg-titan/solon-corpus/claude-projects/AMG_SEO___Social_Content_Competitor_Analysis_Proposal_Builde_019d1d6a/docs.json
  kb:alexander:hormozi     /opt/amg-docs/external-kbs/hormozi/  (.docx + .pdf)
  kb:alexander:welby       /opt/amg-docs/external-kbs/welby/    (.docx)
  kb:alexander:koray       /opt/amg-docs/external-kbs/koray/    (.docx + .txt — skip MP3)
  kb:alexander:reputation  /opt/amg-titan/solon-corpus/claude-projects/AMG_SHIELD___Reputation_Management___Online_Presence_Strateg_019d218f/docs.json
  kb:alexander:paid-ads    /opt/amg-titan/solon-corpus/claude-projects/AMG_Paid_Ads_Strategist_v1_0_019d3acd/docs.json
  kb:nestor:lumina-cro     /opt/amg-titan/solon-corpus/claude-projects/AMG_CRO___CONVERSION_DESIGN_STRATEGIST__Lumina__019d1152/docs.json

10 namespaces total.

Vector store schema (op_memory_vectors):
  id            uuid (deterministic from source_path:chunk_index sha256 → uuid)
  content       text (the chunk)
  summary       text (first 180 chars)
  embedding     vector(768) (nomic-embed-text)
  project_tag   text (the chief — 'hercules', 'alexander', 'nestor')
  project_id    text (same)
  chunk_type    'kb'
  operator_id   'KB_INGEST_PHASE_B'
  model_name    'nomic-embed-text'
  embedding_dim 768
  topic_tags    [namespace, source_filename, project_uuid_if_any]
  status        'active'
  pinned        false
  muted         false

Usage (on VPS):
  python3 phase_b_kb_ingest.py --inventory          # dry run, print sources + counts
  python3 phase_b_kb_ingest.py --namespace kb:hercules:eom  # ingest one namespace
  python3 phase_b_kb_ingest.py --all                # ingest all 10
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import pathlib
import re
import subprocess
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
from datetime import datetime, timezone

OLLAMA_URL = os.environ.get("OLLAMA_URL", "http://127.0.0.1:11434")
EMBED_MODEL = os.environ.get("EMBED_MODEL", "nomic-embed-text")
SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://egoazyasyrhslluossli.supabase.co")
SUPABASE_KEY = (os.environ.get("SUPABASE_SERVICE_ROLE_KEY")
                or os.environ.get("SUPABASE_AMG_PROD_SERVICE_ROLE_KEY")
                or "")

# Chunking parameters per the briefing.
CHUNK_TOKENS = 500
OVERLAP_TOKENS = 50
# Rough word count per token = 0.75; use that to convert tokens → chars.
# 500 tokens ≈ 1875 chars; 50 token overlap ≈ 188 chars.
CHUNK_CHARS = int(CHUNK_TOKENS * 4)  # 2000 chars
OVERLAP_CHARS = int(OVERLAP_TOKENS * 4)  # 200 chars

LOG_FILE = pathlib.Path("/var/log/phase_b_kb_ingest.log")

# Namespace → (chief, source_kind, source_paths)
SOURCE_KIND_DOCS_JSON = "docs_json"
SOURCE_KIND_DIR_OF_FILES = "dir_of_files"
SOURCE_KIND_MERGED_DIRS = "merged_dirs"

NAMESPACE_PLAN = [
    {
        "namespace": "kb:hercules:eom",
        "chief": "hercules",
        "source_kind": SOURCE_KIND_DOCS_JSON,
        "sources": ["/opt/amg-titan/solon-corpus/claude-projects/AMG_Executive_Operations_Manager_v2_0_019d3a46/docs.json"],
    },
    {
        "namespace": "kb:hercules:doctrine",
        "chief": "hercules",
        "source_kind": SOURCE_KIND_MERGED_DIRS,
        "sources": ["/opt/amg-docs/doctrines/", "/opt/amg-docs/doctrine/"],
        "extensions": [".md", ".txt"],
    },
    {
        "namespace": "kb:alexander:outbound",
        "chief": "alexander",
        "source_kind": SOURCE_KIND_DOCS_JSON,
        "sources": ["/opt/amg-titan/solon-corpus/claude-projects/AMG_Outbound_Leadgen_Advisor_019d1d70/docs.json"],
    },
    {
        "namespace": "kb:alexander:seo-content",
        "chief": "alexander",
        "source_kind": SOURCE_KIND_DOCS_JSON,
        "sources": ["/opt/amg-titan/solon-corpus/claude-projects/AMG_SEO___Social_Content_Competitor_Analysis_Proposal_Builde_019d1d6a/docs.json"],
    },
    {
        "namespace": "kb:alexander:hormozi",
        "chief": "alexander",
        "source_kind": SOURCE_KIND_DIR_OF_FILES,
        "sources": ["/opt/amg-docs/external-kbs/hormozi/"],
        "extensions": [".docx", ".pdf"],
    },
    {
        "namespace": "kb:alexander:welby",
        "chief": "alexander",
        "source_kind": SOURCE_KIND_DIR_OF_FILES,
        "sources": ["/opt/amg-docs/external-kbs/welby/"],
        "extensions": [".docx"],
    },
    {
        "namespace": "kb:alexander:koray",
        "chief": "alexander",
        "source_kind": SOURCE_KIND_DIR_OF_FILES,
        "sources": ["/opt/amg-docs/external-kbs/koray/"],
        "extensions": [".docx", ".txt"],
    },
    {
        "namespace": "kb:alexander:reputation",
        "chief": "alexander",
        "source_kind": SOURCE_KIND_DOCS_JSON,
        "sources": ["/opt/amg-titan/solon-corpus/claude-projects/AMG_SHIELD___Reputation_Management___Online_Presence_Strateg_019d218f/docs.json"],
    },
    {
        "namespace": "kb:alexander:paid-ads",
        "chief": "alexander",
        "source_kind": SOURCE_KIND_DOCS_JSON,
        "sources": ["/opt/amg-titan/solon-corpus/claude-projects/AMG_Paid_Ads_Strategist_v1_0_019d3acd/docs.json"],
    },
    {
        "namespace": "kb:nestor:lumina-cro",
        "chief": "nestor",
        "source_kind": SOURCE_KIND_DOCS_JSON,
        "sources": ["/opt/amg-titan/solon-corpus/claude-projects/AMG_CRO___CONVERSION_DESIGN_STRATEGIST__Lumina__019d1152/docs.json"],
    },
]


def _log(msg: str) -> None:
    LOG_FILE.parent.mkdir(parents=True, exist_ok=True)
    ts = datetime.now(timezone.utc).isoformat()
    line = f"[{ts}] {msg}"
    try:
        with LOG_FILE.open("a") as f:
            f.write(line + "\n")
    except Exception:
        pass
    print(line, file=sys.stderr)


# ─── parsers ────────────────────────────────────────────────────────────────
def _parse_docx(path: pathlib.Path) -> str:
    try:
        from docx import Document  # python-docx
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        _log(f"docx parse FAILED {path}: {e!r}")
        return ""


def _parse_pdf(path: pathlib.Path) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    text_parts.append(t)
        return "\n\n".join(text_parts)
    except Exception as e:
        _log(f"pdfplumber FAILED {path}: {e!r}; trying pandoc")
    # Fallback: pandoc
    try:
        out = subprocess.run(["pandoc", "-t", "plain", str(path)],
                             capture_output=True, text=True, timeout=60)
        if out.returncode == 0 and out.stdout.strip():
            return out.stdout
    except Exception as e:
        _log(f"pandoc fallback FAILED {path}: {e!r}")
    return ""


def _parse_text_file(path: pathlib.Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        _log(f"text read FAILED {path}: {e!r}")
        return ""


def _extract_text(path: pathlib.Path) -> str:
    suf = path.suffix.lower()
    if suf == ".docx":
        return _parse_docx(path)
    if suf == ".pdf":
        return _parse_pdf(path)
    if suf in {".md", ".txt"}:
        return _parse_text_file(path)
    return ""


# ─── chunking ───────────────────────────────────────────────────────────────
def chunk_text(text: str, chunk_chars: int = CHUNK_CHARS,
               overlap_chars: int = OVERLAP_CHARS) -> list[str]:
    """Sliding-window chunker. Splits on paragraph boundaries when possible to
    keep semantic coherence. ~500 tokens (2000 chars) per chunk with 50-token
    (200-char) overlap."""
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= chunk_chars:
        return [text]
    out: list[str] = []
    paragraphs = re.split(r"\n\s*\n", text)
    buf = ""
    for para in paragraphs:
        if not para.strip():
            continue
        if len(buf) + len(para) + 2 <= chunk_chars:
            buf = (buf + "\n\n" + para) if buf else para
        else:
            if buf:
                out.append(buf)
            # If a single paragraph is bigger than chunk_chars, hard-split it.
            if len(para) > chunk_chars:
                step = chunk_chars - overlap_chars
                for i in range(0, len(para), step):
                    out.append(para[i:i + chunk_chars])
                buf = ""
            else:
                buf = para
    if buf:
        out.append(buf)

    # Add overlap: prepend the last `overlap_chars` of each chunk to the next.
    overlapped: list[str] = []
    for i, c in enumerate(out):
        if i == 0:
            overlapped.append(c)
        else:
            tail = out[i - 1][-overlap_chars:]
            overlapped.append(tail + "\n\n" + c)
    return overlapped


# ─── ollama embed ───────────────────────────────────────────────────────────
def embed_text(text: str) -> list[float] | None:
    body = {"model": EMBED_MODEL, "prompt": text[:8000]}
    req = urllib.request.Request(
        f"{OLLAMA_URL.rstrip('/')}/api/embeddings",
        data=json.dumps(body).encode(),
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as r:
            data = json.loads(r.read())
        emb = data.get("embedding")
        if not emb or len(emb) != 768:
            _log(f"embed unexpected shape: len={len(emb) if emb else 0}")
            return None
        return emb
    except Exception as e:
        _log(f"embed FAILED: {e!r}")
        return None


# ─── supabase write ─────────────────────────────────────────────────────────
def supabase_upsert_vector(row: dict) -> bool:
    """POST to op_memory_vectors via Supabase REST API. Uses upsert on id."""
    url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/op_memory_vectors"
    headers = {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json",
        "Prefer": "resolution=merge-duplicates,return=minimal",
    }
    req = urllib.request.Request(
        url, data=json.dumps(row).encode(), headers=headers, method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            return 200 <= r.status < 300
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")[:500]
        _log(f"supabase upsert FAILED HTTP {e.code}: {body}")
        return False
    except Exception as e:
        _log(f"supabase upsert exception: {e!r}")
        return False


def _deterministic_uuid(seed: str) -> str:
    h = hashlib.sha256(seed.encode("utf-8")).hexdigest()
    # Build a UUIDv4-shaped string from the hash bytes.
    return f"{h[0:8]}-{h[8:12]}-4{h[13:16]}-a{h[17:20]}-{h[20:32]}"


# ─── source readers ─────────────────────────────────────────────────────────
def _read_docs_json(json_path: pathlib.Path) -> list[dict]:
    """Each Claude Project's docs.json is a JSON array of {uuid, file_name,
    content, created_at, project_uuid, estimated_token_count}. Return as
    {file_name, content, source_uuid}."""
    try:
        data = json.loads(json_path.read_text(encoding="utf-8", errors="replace"))
    except Exception as e:
        _log(f"docs.json read FAILED {json_path}: {e!r}")
        return []
    if not isinstance(data, list):
        return []
    out = []
    for d in data:
        if not isinstance(d, dict):
            continue
        content = d.get("content") or ""
        if not content.strip():
            continue
        out.append({
            "file_name": d.get("file_name") or d.get("uuid") or "unknown",
            "content": content,
            "source_uuid": d.get("uuid") or "",
            "project_uuid": d.get("project_uuid") or "",
        })
    return out


def _read_dir_files(dir_path: pathlib.Path, extensions: list[str]) -> list[dict]:
    """Walk a directory, parse each file by extension, return {file_name,
    content, source_uuid (sha256 of content)}."""
    out = []
    if not dir_path.exists():
        _log(f"dir does not exist: {dir_path}")
        return []
    for p in sorted(dir_path.rglob("*")):
        if not p.is_file():
            continue
        if p.suffix.lower() not in extensions:
            continue
        text = _extract_text(p)
        if not text.strip():
            _log(f"  skip empty: {p}")
            continue
        out.append({
            "file_name": p.name,
            "content": text,
            "source_uuid": hashlib.sha256(text.encode()).hexdigest()[:16],
            "source_path": str(p),
        })
    return out


def _read_merged_dirs(dirs: list[str], extensions: list[str]) -> list[dict]:
    """For doctrine merge: read both /opt/amg-docs/doctrine/ + /opt/amg-docs/doctrines/,
    dedupe by content sha256."""
    seen: dict[str, dict] = {}
    for d in dirs:
        for entry in _read_dir_files(pathlib.Path(d), extensions):
            sha = hashlib.sha256(entry["content"].encode()).hexdigest()
            if sha not in seen:
                seen[sha] = entry
    return list(seen.values())


def read_sources(plan: dict) -> list[dict]:
    kind = plan["source_kind"]
    if kind == SOURCE_KIND_DOCS_JSON:
        out = []
        for s in plan["sources"]:
            out.extend(_read_docs_json(pathlib.Path(s)))
        return out
    if kind == SOURCE_KIND_DIR_OF_FILES:
        out = []
        for s in plan["sources"]:
            out.extend(_read_dir_files(pathlib.Path(s), plan["extensions"]))
        return out
    if kind == SOURCE_KIND_MERGED_DIRS:
        return _read_merged_dirs(plan["sources"], plan["extensions"])
    return []


# ─── orchestration ──────────────────────────────────────────────────────────
def ingest_namespace(plan: dict, dry_run: bool = False) -> dict:
    namespace = plan["namespace"]
    chief = plan["chief"]
    _log(f"=== {namespace} (chief={chief}, kind={plan['source_kind']}) ===")
    sources = read_sources(plan)
    total_chunks = 0
    total_inserted = 0
    total_failed = 0
    file_count = len(sources)
    _log(f"  source files: {file_count}")
    if dry_run:
        for src in sources[:5]:
            _log(f"    sample: {src['file_name']} ({len(src['content'])} chars)")
        return {"namespace": namespace, "files": file_count, "dry_run": True}

    for src in sources:
        chunks = chunk_text(src["content"])
        for i, chunk in enumerate(chunks):
            total_chunks += 1
            seed = f"{namespace}|{src.get('source_uuid','')}|{src['file_name']}|{i}"
            row_id = _deterministic_uuid(seed)
            embedding = embed_text(chunk)
            if embedding is None:
                total_failed += 1
                continue
            row = {
                "id": row_id,
                "content": chunk,
                "summary": chunk[:180],
                "embedding": embedding,
                "project_tag": chief,
                "project_id": chief,
                "chunk_type": "kb",
                "operator_id": "KB_INGEST_PHASE_B",
                "model_name": EMBED_MODEL,
                "embedding_dim": 768,
                "topic_tags": [namespace, src["file_name"], f"chief:{chief}"],
                "status": "active",
                "pinned": False,
                "muted": False,
            }
            ok = supabase_upsert_vector(row)
            if ok:
                total_inserted += 1
            else:
                total_failed += 1
            if total_chunks % 25 == 0:
                _log(f"  progress: {total_inserted}/{total_chunks} inserted")
    _log(f"  DONE {namespace}: files={file_count} chunks={total_chunks} "
         f"inserted={total_inserted} failed={total_failed}")
    return {
        "namespace": namespace, "files": file_count,
        "chunks": total_chunks, "inserted": total_inserted, "failed": total_failed,
    }


def main() -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--inventory", action="store_true",
                   help="dry run — print sources + counts, no embedding/insert")
    p.add_argument("--namespace", type=str, default=None,
                   help="ingest one namespace only (e.g. kb:hercules:eom)")
    p.add_argument("--all", action="store_true",
                   help="ingest all 10 namespaces sequentially")
    args = p.parse_args()

    if not SUPABASE_KEY and not args.inventory:
        print("error: SUPABASE_SERVICE_ROLE_KEY (or SUPABASE_AMG_PROD_SERVICE_ROLE_KEY) "
              "must be set in env", file=sys.stderr)
        return 1

    selected = NAMESPACE_PLAN
    if args.namespace:
        selected = [p for p in NAMESPACE_PLAN if p["namespace"] == args.namespace]
        if not selected:
            print(f"unknown namespace: {args.namespace}", file=sys.stderr)
            return 1

    results = []
    for plan in selected:
        try:
            r = ingest_namespace(plan, dry_run=args.inventory)
        except Exception as e:
            _log(f"  EXCEPTION {plan['namespace']}: {e!r}")
            r = {"namespace": plan["namespace"], "error": repr(e)}
        results.append(r)

    print("\n=== INGESTION SUMMARY ===")
    print(json.dumps(results, indent=2, default=str))
    fails = sum(r.get("failed", 0) for r in results)
    inserted = sum(r.get("inserted", 0) for r in results)
    print(f"\nTotal inserted: {inserted}, failed: {fails}")
    return 0 if not fails else 1


if __name__ == "__main__":
    sys.exit(main())

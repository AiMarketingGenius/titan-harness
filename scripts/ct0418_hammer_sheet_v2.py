#!/usr/bin/env python3
"""
titan-harness/scripts/ct0418_hammer_sheet_v2.py

HAMMER_SHEET v2.0 — canonical Solon-voice phrase library for the Don Martelli
pitch and every future Chamber AI Advantage conversation.

Per Solon directive 2026-04-18T19:40Z (Hammer Sheet v2.0 addendum).

Six categories on one mobile-readable doc:
  1. OPENING FRAMES  (Memory Guard hook + Microsoft analogy)
  2. ECONOMIC MATH   (tiered rev-share, 35% total margin, retained discount, dues 100-500%)
  3. CIVIC SOVEREIGNTY CLOSERS ("City Hall begging YOU", economic engine of cities)
  4. CO-FOUNDER LANGUAGE  (shared legacy, Premier Partner unlock, Advisory Board seat #1)
  5. GUARANTEE  (verbatim locked)
  6. OBJECTION COUNTERS  (Don-specific likely questions + counters)

Output:
  plans/deployments/HAMMER_SHEET_v2_0_2026-04-18.docx
  VPS mirror: /opt/amg-outbound/pitch-prep/HAMMER_SHEET_v2_0_2026-04-18.docx (separate scp)

Ship tag: hammer-sheet-v2-0-2026-04-18

Run: python3 scripts/ct0418_hammer_sheet_v2.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = REPO_ROOT / "plans" / "deployments"

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
TEAL = RGBColor(0x26, 0xC6, 0xA8)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
AMBER = RGBColor(0xC9, 0x7A, 0x17)
VIOLET = RGBColor(0x7E, 0x57, 0xC2)
TEXT = RGBColor(0x1C, 0x1C, 0x1E)
MUTED = RGBColor(0x5F, 0x6A, 0x72)


def _style_normal(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT


def _h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY


def _h2(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = color


def _h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY


def _quote(doc, text, color=TEAL):
    """Verbatim Solon phrase — italic, colored, indented."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(12)
    r = p.add_run(f"\u201C{text}\u201D")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = color


def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(11)


def _meta(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def build_hammer_sheet() -> Path:
    doc = Document()
    _style_normal(doc)

    _h1(doc, "HAMMER SHEET v2.1")
    _meta(doc, "Chamber AI Advantage · Solon voice canonical phrase library · 2026-04-18")
    _meta(doc, "v2.1 addendum: Chamber-Only Distribution doctrine + Gatekeeper reframe per Solon 20:10Z.")
    _meta(doc, "Internalize before Monday 2026-04-20 Don Martelli pitch. Memorized lines beat scripted ones.")

    _body(
        doc,
        "How to read this: every line in a colored quote block is a verbatim Solon phrase. "
        "Memorize the shape and the rhythm. The goal isn't to recite — it's to have the line "
        "ready the moment the opening for it appears. When Don says X, you reach for Y. That's "
        "the whole Hammer Sheet.",
    )

    # ─── 1. OPENING FRAMES ──────────────────────────────────
    _h2(doc, "1. Opening Frames", TEAL)
    _body(doc, "Two opens, in this order — the Memory Guard hook and the Microsoft analogy. Together they land the AI-is-broken problem and the only-Chambers-can-sell-the-fix exclusivity in under 90 seconds.")

    _h3(doc, "Memory Guard opening (set the problem)")
    _quote(doc,
        "Every AI tool forgets, hallucinates, and drifts. We built the first one that "
        "remembers — verified, across every LLM you use, with real fact-checking and drift "
        "protection. Watch.")

    _h3(doc, "Microsoft analogy (set the exclusivity moat)")
    _quote(doc,
        "Imagine if Microsoft Windows was only sold through local Chambers of Commerce. "
        "That's Chamber AI Advantage. Every small business on Earth needs AI marketing in "
        "2026. Every local business already trusts its Chamber. This is the first product "
        "ever built to be sold through only one channel — and if Revere signs, that channel "
        "is yours.")

    _h3(doc, "Read-back to check landing")
    _quote(doc, "Does that land? Because everything downstream depends on you buying that frame.")

    # ─── 2. ECONOMIC MATH ───────────────────────────────────
    _h2(doc, "2. Economic Math", GREEN)

    _h3(doc, "Tiered rev-share table (verbatim)")
    _bullet(doc, "1 – 5 subs: 10% — on-ramp tier. $339/mo at 5 Growth subs.")
    _bullet(doc, "6 – 10 subs: 15% — $1,016/mo at 10 Growth subs.")
    _bullet(doc, "11 – 20 subs: 17% — $2,302/mo at 20 Growth subs.")
    _bullet(doc, "21 – 30 subs: 19% — $3,859/mo at 30 Growth subs.")
    _bullet(doc, "30+ subs: 20% + Premier Partner unlock — $6,770/mo at 50 Growth subs.")

    _h3(doc, "35% total margin framing")
    _quote(doc,
        "Stack the rev-share on top of the retained-discount lever, cap out past thirty "
        "active subs, and each Growth-tier member generates up to two hundred seventy-nine "
        "dollars a month for the Chamber. Fifty members at that economics is fourteen "
        "thousand a month. That's a real budget line for a real Chamber Foundation.")

    _h3(doc, "Retained-discount mechanic (one-liner)")
    _quote(doc,
        "You pick the path. Pass the 15% member discount to your members, or keep the "
        "discount as Chamber margin and have members pay retail. Same seven agents either "
        "way. Your choice per cohort, or Chamber-wide.")

    _h3(doc, "Dues compounding (100-500% lift)")
    _quote(doc,
        "The rev-share isn't even the biggest number. When your members tell two other "
        "local businesses \u2018my Chamber handed me an AI marketing team,\u2019 membership "
        "stops being a line item and becomes the single best marketing decision any local "
        "business makes. Dues revenue typically lifts 100 to 500 percent within two years. "
        "The dues lift compounds on top of the rev-share, not instead of it.")

    # ─── 3. CIVIC SOVEREIGNTY CLOSERS ────────────────────────
    _h2(doc, "3. Civic Sovereignty Closers", AMBER)

    _h3(doc, "Economic engine framing (anchor closer)")
    _quote(doc,
        "Chambers become the economic engine of cities, not ribbon-cutting committees.",
        color=AMBER)

    _h3(doc, "City Hall begs YOU (the high-emotion closer)")
    _quote(doc,
        "The endgame: City Hall begs YOU to help them fund the next downtown project — "
        "not the other way around.",
        color=AMBER)

    _h3(doc, "Civic Legacy bullet list (what the money funds)")
    _bullet(doc, "Give grants to member businesses hit by fire, flood, or downturn.")
    _bullet(doc, "Fund startup loans for new local entrepreneurs — zero-interest, underwritten by the Chamber.")
    _bullet(doc, "Host charity galas the Chamber pays for, not chases sponsors for.")
    _bullet(doc, "Create jobs by seeding member businesses' expansion into new locations.")
    _bullet(doc, "Drive economic prosperity with downtown façade grants + small-business mentorship.")
    _bullet(doc, "Donate to causes the Chamber board cares about — veterans, literacy, youth sports.")
    _bullet(doc, "Fund scholarships for the children of member employees.")
    _bullet(doc, "Stand up a 501(c) Chamber Foundation with sustainable recurring revenue.")

    _h3(doc, "The \u201Cyou stop being\u201D reframe")
    _quote(doc,
        "You stop being a ribbon-cutting committee. You become the economic engine of Revere. "
        "You're not cashing rev-share checks — you're rebuilding Revere.")

    # ─── 3.5 CHAMBER-ONLY DISTRIBUTION · GATEKEEPER FRAMING ──
    _h2(doc, "3.5. Chamber-Gatekeeper Reframe (v2.1 addendum · Solon 20:10Z)", AMBER)
    _body(doc, "The exclusivity moat isn't only that Chambers can sell AMG. It's that Chambers are the ONLY way ANY business — SMB, mid-market, enterprise — buys AMG products in that Chamber's region. Every AMG product. Every price point. Sold only through Chambers. Use these phrases when the conversation moves from \"member services\" to \"what this scales into.\"")

    _h3(doc, "The core line (use verbatim)")
    _quote(doc,
        "Every AMG product. SMB, mid-sized, enterprise Atlas deployments. Sold only "
        "through Chambers. No exceptions. Corporate America pays. Main Street receives.",
        color=AMBER)

    _h3(doc, "The Don-specific pitch")
    _quote(doc,
        "Don, when State Street or CVS eventually want our Atlas system, they don't talk "
        "to me. They talk to you. Your Chamber becomes the gatekeeper every business in "
        "your region must go through. That's the power we're handing you — not a software "
        "reseller agreement. A civic gatekeeper role.",
        color=AMBER)

    _h3(doc, "The counter to \"will big businesses accept that?\"")
    _quote(doc,
        "Solon pilots every enterprise executive-board relationship personally — I'm in "
        "the room for the State-Street-sized calls, not the Chamber. The Chamber is the "
        "merchant of record. Customer experience is identical to direct engagement. "
        "Only difference is payment routing: their check clears through the Chamber "
        "instead of direct to AMG. Corporate America pays. Main Street receives.",
        color=AMBER)

    _h3(doc, "The permanence line")
    _quote(doc,
        "This isn't a tier. This isn't a bonus. This is the doctrine. Encyclopedia v1.4.2 "
        "Section 30 locks it: AMG never sells direct, any product, any customer size, "
        "anywhere a Chamber partner exists. Your Chamber isn't one of many channels. "
        "Your Chamber is THE channel.",
        color=AMBER)

    # ─── 4. CO-FOUNDER LANGUAGE ──────────────────────────────
    _h2(doc, "4. Co-Founder Language", VIOLET)

    _h3(doc, "Shared legacy opener (lead-with talking point #1)")
    _quote(doc,
        "Don, I'm not pitching you as a customer. I'm pitching you as the co-founder of "
        "what this becomes nationally. Your name first in every future Founding Partner "
        "conversation. Your Chamber as the validation case study for every pitch nationwide. "
        "This is shared legacy, not sponsorship.",
        color=VIOLET)

    _h3(doc, "Your name first framing")
    _quote(doc,
        "Five years from now, when Chamber AI Advantage is running in three hundred Chambers "
        "across the country, that first Chamber that signed is the one everybody references. "
        "The co-architect position is only available to the first Chamber that signs. I'd "
        "rather it be Revere.",
        color=VIOLET)

    _h3(doc, "Advisory Board seat #1 (concrete lever)")
    _quote(doc,
        "Advisory Board seat number one, on the record. Voting voice on program direction, "
        "pricing evolution, Chamber-facing feature priorities. Right of first look on every "
        "new Chamber AI Advantage product line thirty days ahead of the general network.",
        color=VIOLET)

    _h3(doc, "Premier Partner Atlas reseller unlock")
    _quote(doc,
        "At thirty active subs, Revere unlocks Premier Partner status — the right to resell "
        "AMG's custom enterprise AI builds, twenty-five to two-fifty-K engagements, to non-"
        "member businesses in your territory. Same tiered rev-share. Oracle-style VAR channel. "
        "Chamber gets the warm lead, AMG delivers, Chamber collects the share.",
        color=VIOLET)

    # ─── 5. GUARANTEE ────────────────────────────────────────
    _h2(doc, "5. Guarantee", GREEN)
    _body(doc, "Verbatim. Locked. No hedging, no shortening, no rewording.")
    _quote(doc,
        "If after 3 months you're not completely satisfied with the results we're getting "
        "for you, we'll work a full month FREE. No asterisks. No hedging.",
        color=GREEN)

    _h3(doc, "Pilot-level guarantee follow-up")
    _quote(doc,
        "This applies to the Chamber-level relationship AND each member-level subscription. "
        "If a pilot member isn't satisfied after 90 days, AMG works month four free for them. "
        "If the Chamber isn't satisfied with the partnership itself, AMG works month four "
        "free for the Chamber.",
        color=GREEN)

    # ─── 6. OBJECTION COUNTERS ──────────────────────────────
    _h2(doc, "6. Objection Counters", NAVY)
    _body(doc, "Don's likely questions + Solon-voice responses. Don't paraphrase — the rhythm is the hammer.")

    _h3(doc, "\"Why me?\" (response to co-founder framing)")
    _quote(doc,
        "Because you're Revere. You run the Chamber small businesses tell each other about — "
        "the one that actually returns calls. You know marketing math from thirty years of "
        "running Martelli. And Revere is the territory shape — urban plus suburban plus ~280 "
        "member mix — that makes the validation case study land for the next hundred Chambers. "
        "This isn't random. It's you specifically.")

    _h3(doc, "\"What's the catch?\"")
    _quote(doc,
        "No catch. Zero cost to the Chamber. Rev-share only on what your members actually "
        "buy. Ninety-day pilot with a money-back-equivalent guarantee. The co-architect status "
        "is real — written into the partnership addendum, not a handshake. The upside is "
        "skewed toward YOU, not AMG. That's on purpose, because we need the first Chamber to "
        "win loudly.")

    _h3(doc, "\"What if I pass?\"")
    _quote(doc,
        "Respect the no. If Revere passes, I go to the next Chamber on my list this week — "
        "probably Salem or Beverly, maybe Lynn. One of them signs. Five years from now, when "
        "Chamber AI Advantage is running in three hundred Chambers across the country, that "
        "Chamber is the one everybody references. That's why I'm here first.")

    _h3(doc, "\"My members are too small for AI marketing.\"")
    _quote(doc,
        "That's exactly why they need this. AI marketing is the only marketing that scales "
        "down. A four-person shop gets a seven-agent team for $422 a month — Chamber member "
        "rate, 15% below public retail. A marketing agency would quote them eight to twelve K "
        "a month. The math works BECAUSE your members are small. They're the exact customer "
        "this was built for.")

    _h3(doc, "\"How do I know this actually works?\"")
    _quote(doc,
        "That's the pilot's job. Five to ten of your members, ninety days, zero cost to the "
        "Chamber or to them. At Day 90 you have numbers from real Revere businesses — not my "
        "case studies, yours. You decide what happens after that.")

    _h3(doc, "\"I need to take this to the board.\"")
    _quote(doc,
        "Good. Let's make that easy. One page, three bullets: (1) zero cost to the Chamber, "
        "(2) tiered rev-share up to 20% plus retained-discount lever that stacks to 35% "
        "margin, funding the Chamber Foundation, (3) three-month guarantee or we work a month "
        "free. I'll have that on your desk tomorrow morning.")

    _h3(doc, "\"What about other Chambers? What if Lynn signs too?\"")
    _quote(doc,
        "Twenty-mile territory protection is baked into the partnership. Lynn can sign their "
        "own Chamber partnership — they can't sign Revere's members. Think McDonald's "
        "franchise: every Chamber in New England can participate, but Revere's member base "
        "is yours for as long as you're an active partner.")

    # ─── CLOSE SEQUENCE ──────────────────────────────────────
    _h2(doc, "Close Sequence · Pick one")

    _h3(doc, "If Don signals yes")
    _quote(doc,
        "Let's get five members named by Friday. I'll build the Chamber-branded landing page "
        "this weekend. Pilot starts the Monday after. Does that timeline work?")

    _h3(doc, "If Don is warm but not yet committed")
    _quote(doc,
        "When can you meet again this week — Thursday or Friday? I'll come back with the "
        "specific pilot member cohort criteria your board would want to see, and we lock "
        "this or walk away with a clear no.")

    # Anti-patterns
    _h2(doc, "Don't-Dos (anti-pattern reminders)")
    _bullet(doc, "Don't use \u201CAI-powered\u201D or \u201Cleverage\u201D or \u201Csynergies\u201D — Don will roll his eyes in 30 seconds.")
    _bullet(doc, "Don't pitch before asking what his members are struggling with. Let him tell you. Then fit the offer.")
    _bullet(doc, "Don't over-explain the agents. Name them, one sentence each, move on. The outcome is the story, not the roster.")
    _bullet(doc, "Don't skip the Microsoft analogy. 30 seconds. It reframes the entire conversation.")
    _bullet(doc, "Don't treat civic-legacy language as soft. It's the emotional anchor — the reason Don signs.")
    _bullet(doc, "Don't leave without a specific follow-up date. \u201CI'll get back to you\u201D is a no.")

    path = OUT_DIR / "HAMMER_SHEET_v2_0_2026-04-18.docx"
    doc.save(str(path))
    return path


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    p = build_hammer_sheet()
    print(f"wrote: {p}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
